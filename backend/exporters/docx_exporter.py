"""
DOCX exporter — converts a blueprint dict into a .docx file using python-docx.
"""
from __future__ import annotations

import io
from typing import Any, Dict, List

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    if level == 1:
        h.runs[0].font.color.rgb = RGBColor(0x0F, 0x43, 0x9C)
    elif level == 2:
        h.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    return h


def _add_kv(doc: Document, key: str, value: str):
    p = doc.add_paragraph()
    run_key = p.add_run(f"{key}: ")
    run_key.bold = True
    run_key.font.size = Pt(11)
    run_val = p.add_run(str(value))
    run_val.font.size = Pt(11)


def _add_bullet_list(doc: Document, items: List[str]):
    for item in items:
        doc.add_paragraph(str(item), style="List Bullet")


def _add_section_divider(doc: Document):
    doc.add_paragraph("")


def _inr(val: Any) -> str:
    try:
        return f"₹{int(val):,}"
    except (ValueError, TypeError):
        return str(val)


# ---------------------------------------------------------------------------
# Main exporter
# ---------------------------------------------------------------------------

def export_docx(blueprint: Dict[str, Any]) -> bytes:
    """Return DOCX bytes for the given blueprint dict."""
    doc = Document()

    # Title
    title = doc.add_heading("Startup Blueprint", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    idea_para = doc.add_paragraph()
    idea_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = idea_para.add_run(f"Idea: {blueprint.get('idea', '')}")
    run.font.size = Pt(13)
    run.italic = True

    doc.add_paragraph("")

    # ── Intent Analysis ────────────────────────────────────────────────────
    _add_heading(doc, "1. Startup Intent Analysis", level=1)
    intent = blueprint.get("intent", {})
    fields = [
        ("Sector", intent.get("sector", "")),
        ("Sub-Sector", intent.get("sub_sector", "")),
        ("Business Type", intent.get("business_type", "")),
        ("Target Market", intent.get("target_market", "")),
        ("Stage", intent.get("stage", "")),
        ("Geography", intent.get("geography", "")),
        ("Problem Statement", intent.get("problem_statement", "")),
        ("Unique Value Proposition", intent.get("unique_value_proposition", "")),
        ("Revenue Model", intent.get("revenue_model", "")),
    ]
    for k, v in fields:
        if v:
            _add_kv(doc, k, v)
    _add_section_divider(doc)

    # ── Business Model Canvas ──────────────────────────────────────────────
    _add_heading(doc, "2. Business Model Canvas", level=1)
    bmc = blueprint.get("business_model_canvas", {})
    bmc_labels = {
        "customer_segments": "Customer Segments",
        "value_propositions": "Value Propositions",
        "channels": "Channels",
        "customer_relationships": "Customer Relationships",
        "revenue_streams": "Revenue Streams",
        "key_resources": "Key Resources",
        "key_activities": "Key Activities",
        "key_partnerships": "Key Partnerships",
        "cost_structure": "Cost Structure",
    }
    for key, label in bmc_labels.items():
        _add_heading(doc, label, level=2)
        val = bmc.get(key, "")
        if isinstance(val, list):
            _add_bullet_list(doc, val)
        else:
            doc.add_paragraph(str(val))
    _add_section_divider(doc)

    # ── Budget Estimate ────────────────────────────────────────────────────
    _add_heading(doc, "3. Budget Estimate", level=1)
    budget = blueprint.get("budget_estimate", {})
    currency = budget.get("currency", "INR")
    doc.add_paragraph(f"⚠ Disclaimer: {budget.get('disclaimer', '')}")
    doc.add_paragraph("")

    _add_heading(doc, "One-Time Setup Costs", level=2)
    for item in budget.get("one_time_costs", []):
        doc.add_paragraph(
            f"• {item['item']}: {_inr(item['min'])} – {_inr(item['max'])}  |  {item.get('notes', '')}",
            style="List Bullet",
        )
    _add_kv(doc, "Total Setup Cost", f"{_inr(budget.get('total_setup_min', 0))} – {_inr(budget.get('total_setup_max', 0))}")

    doc.add_paragraph("")
    _add_heading(doc, "Monthly Recurring Costs", level=2)
    for item in budget.get("monthly_recurring", []):
        doc.add_paragraph(
            f"• {item['item']}: {_inr(item['min'])} – {_inr(item['max'])}/month  |  {item.get('notes', '')}",
            style="List Bullet",
        )
    _add_kv(doc, "Monthly Burn Rate", f"{_inr(budget.get('monthly_burn_min', 0))} – {_inr(budget.get('monthly_burn_max', 0))}/month")
    _add_kv(doc, "Recommended Runway", budget.get("runway_recommendation", "12–18 months"))
    _add_kv(doc, "Estimated Funding Needed (12 months)", f"{_inr(budget.get('funding_needed_min', 0))} – {_inr(budget.get('funding_needed_max', 0))}")
    _add_section_divider(doc)

    # ── GTM Strategy ───────────────────────────────────────────────────────
    _add_heading(doc, "4. Go-To-Market Strategy", level=1)
    gtm = blueprint.get("gtm_strategy", {})
    _add_kv(doc, "Target Beachhead", gtm.get("target_beachhead", ""))
    _add_kv(doc, "Positioning Statement", gtm.get("positioning_statement", ""))
    doc.add_paragraph("")
    _add_heading(doc, "Key Success Metrics", level=2)
    _add_bullet_list(doc, gtm.get("key_metrics", []))
    _add_kv(doc, "CAC Estimate", gtm.get("cac_estimate", ""))
    _add_kv(doc, "LTV Estimate", gtm.get("ltv_estimate", ""))
    doc.add_paragraph("")

    for phase in gtm.get("phases", []):
        _add_heading(doc, phase.get("phase", ""), level=2)
        _add_kv(doc, "Duration", phase.get("duration", ""))
        _add_kv(doc, "Goal", phase.get("goal", ""))
        _add_kv(doc, "Channels", ", ".join(phase.get("channels", [])))
        _add_heading(doc, "Tactics", level=3)
        _add_bullet_list(doc, phase.get("tactics", []))
        _add_kv(doc, "Success Criteria", phase.get("success_criteria", ""))
        doc.add_paragraph("")
    _add_section_divider(doc)

    # ── Competitor Analysis ────────────────────────────────────────────────
    _add_heading(doc, "5. Competitor Analysis", level=1)
    for i, comp in enumerate(blueprint.get("competitors", []), 1):
        _add_heading(doc, f"{i}. {comp.get('name', '')}", level=2)
        _add_kv(doc, "Sector", comp.get("sector", ""))
        _add_kv(doc, "Business Model", comp.get("business_model", ""))
        doc.add_paragraph(comp.get("description", ""))
        _add_heading(doc, "Strengths", level=3)
        _add_bullet_list(doc, comp.get("strengths", []))
        _add_heading(doc, "Weaknesses", level=3)
        _add_bullet_list(doc, comp.get("weaknesses", []))
        _add_kv(doc, "Differentiation Opportunity", comp.get("positioning_gap", ""))
        if comp.get("source"):
            _add_kv(doc, "Source", comp.get("source", ""))
        doc.add_paragraph("")
    _add_section_divider(doc)

    # ── Government Schemes ─────────────────────────────────────────────────
    _add_heading(doc, "6. Applicable Government Schemes", level=1)
    for scheme in blueprint.get("government_schemes", []):
        _add_heading(doc, scheme.get("name", ""), level=2)
        _add_kv(doc, "Provider", scheme.get("provider", ""))
        doc.add_paragraph(scheme.get("description", ""))
        _add_kv(doc, "Funding Amount", scheme.get("funding_amount", ""))
        _add_kv(doc, "Applicable Stages", ", ".join(scheme.get("stage", [])))
        _add_heading(doc, "Eligibility Criteria", level=3)
        _add_bullet_list(doc, scheme.get("eligibility", []))
        if scheme.get("source"):
            _add_kv(doc, "Source", scheme.get("source", ""))
        if scheme.get("last_updated"):
            _add_kv(doc, "Data Last Updated", scheme.get("last_updated", ""))
        doc.add_paragraph("")
    _add_section_divider(doc)

    # ── Investors ──────────────────────────────────────────────────────────
    _add_heading(doc, "7. Investor & Incubator Matches", level=1)
    for inv in blueprint.get("investors", []):
        _add_heading(doc, inv.get("name", ""), level=2)
        _add_kv(doc, "Type", inv.get("type", ""))
        _add_kv(doc, "Ticket Size", inv.get("ticket_size", ""))
        _add_kv(doc, "Stages", ", ".join(inv.get("stages", [])))
        _add_kv(doc, "Focus Sectors", ", ".join(inv.get("focus_sectors", [])[:5]))
        doc.add_paragraph(inv.get("description", ""))
        if inv.get("application"):
            _add_kv(doc, "Apply At", inv.get("application", ""))
        doc.add_paragraph("")
    _add_section_divider(doc)

    # ── Legal & Compliance ─────────────────────────────────────────────────
    _add_heading(doc, "8. Legal & Compliance Checklist", level=1)
    legal = blueprint.get("legal_compliance", {})
    _add_kv(doc, "Recommended Structure", legal.get("recommended_structure", ""))
    _add_kv(doc, "Reason", legal.get("reason", ""))
    doc.add_paragraph("")

    _add_heading(doc, "Registration Steps", level=2)
    for step in legal.get("registration_steps", []):
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{step['step']}")
        r.bold = True
        p.add_run(f" — {step['authority']} | Timeline: {step['timeline']} | Cost: {step['cost_approx']} | {step['priority'].upper()}")

    doc.add_paragraph("")
    _add_heading(doc, "Ongoing Compliance", level=2)
    for item in legal.get("compliance_checklist", []):
        doc.add_paragraph(
            f"{item['item']} ({item['frequency']}) — {item['authority']}. {item.get('notes', '')}",
            style="List Bullet",
        )

    doc.add_paragraph("")
    _add_heading(doc, "IP Recommendations", level=2)
    _add_bullet_list(doc, legal.get("ip_recommendations", []))

    doc.add_paragraph("")
    _add_heading(doc, "Key Legal Risks", level=2)
    _add_bullet_list(doc, legal.get("key_risks", []))
    _add_section_divider(doc)

    # ── Sources ────────────────────────────────────────────────────────────
    sources = blueprint.get("sources", [])
    if sources:
        _add_heading(doc, "Sources & References", level=1)
        for src in sources:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{src.get('name', '')} ({src.get('type', '')})").bold = True
            if src.get("url"):
                p.add_run(f" — {src.get('url', '')}")

    # Footer note
    doc.add_paragraph("")
    footer_para = doc.add_paragraph(
        "Generated by Startup Blueprint Generator | Powered by IBM Granite + IBM Watson x.ai | Data grounded in verified knowledge base"
    )
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.runs[0].font.size = Pt(9)
    footer_para.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # Serialise to bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
